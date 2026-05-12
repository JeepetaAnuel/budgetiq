from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy
import re

doc = Document()

# ─── Page setup ───
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# ─── Styles ───
style = doc.styles['Normal']
font = style.font
font.name = 'Inter'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.4

TEAL = RGBColor(0x0d, 0x94, 0x88)
DARK = RGBColor(0x0f, 0x17, 0x2a)
GRAY = RGBColor(0x47, 0x55, 0x69)
MUTED = RGBColor(0x94, 0xa3, 0xb8)
WHITE = RGBColor(0xff, 0xff, 0xff)
GREEN = RGBColor(0x10, 0xb9, 0x81)
RED = RGBColor(0xef, 0x44, 0x44)
PURPLE = RGBColor(0x63, 0x66, 0xf1)
BG_DARK = RGBColor(0x0f, 0x17, 0x2a)

def add_colored_box(doc, color, text, bold=False, size=9, text_color=WHITE):
    """Add a shaded paragraph (simulates a card)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = text_color
    # Add shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    p._element.get_or_add_pPr().append(shading)
    return p

def add_tag(doc, text, color=TEAL, bg="E8F5F0"):
    """Inline tag"""
    p = doc.add_paragraph()
    run = p.add_run(f"  {text}  ")
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = color
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
    p.paragraph_format._element.get_or_add_pPr().append(shading)
    return p

def add_section_title(doc, icon, title, tag_text, tag_color):
    """Section header with icon, title, and tag"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    border = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="E2E8F0"/>'
        f'</w:pBdr>'
    )
    p.paragraph_format._element.get_or_add_pPr().append(border)
    run = p.add_run(f"{icon}  {title}    ")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = DARK
    run2 = p.add_run(tag_text)
    run2.font.size = Pt(8)
    run2.font.bold = True
    run2.font.color.rgb = tag_color
    return p

def add_bullet(doc, text, bold_part=""):
    """Add a bullet point with optional bold prefix"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    if bold_part:
        r = p.add_run(bold_part)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = DARK
        r2 = p.add_run(text)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = GRAY
    else:
        r = p.add_run(text)
        r.font.size = Pt(9.5)
        r.font.color.rgb = GRAY
    return p

def add_phone_mockup_table(doc, rows_data, bg_color="F8FAFC"):
    """Creates a phone-mockup-like bordered table"""
    rows = len(rows_data)
    table = doc.add_table(rows=rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value, value_color) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(2.5)
        # Label
        lp = row.cells[0].paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lr = lp.add_run(label)
        lr.font.size = Pt(8)
        lr.font.color.rgb = MUTED
        # Value
        vp = row.cells[1].paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        vr = vp.add_run(str(value))
        vr.font.size = Pt(9)
        vr.font.bold = True
        vr.font.color.rgb = value_color if value_color else DARK
        # Shading
        fill_val = bg_color if i % 2 == 0 else "FFFFFF"
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_val}"/>')
        row.cells[0].paragraphs[0]._element.get_or_add_pPr().append(deepcopy(shading))
        row.cells[1].paragraphs[0]._element.get_or_add_pPr().append(deepcopy(shading))
    # Remove table borders and set cell padding
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

FEATURES = [
    ("🏠", "Dashboard / Inicio", "Implementado", TEAL, [
        "Saldo total con gráfico de tendencia de 14 días",
        "Ingresos / Gastos / Ahorro del mes actual",
        "Health Score financiero (0-100) con color según estado",
        "Acciones rápidas: Modo Manual, Modo IA, Presupuestos, Escanear",
        "Actividad reciente con últimas 6 transacciones",
        "Insights automáticos según comportamiento de gasto",
    ], [
        ("Saldo · Este mes", "$2,847.50", None),
        ("Ingresos", "+$3,200.00", GREEN),
        ("Gastos", "-$352.50", RED),
        ("Ahorro", "+$2,847.50", GREEN),
        ("Health Score", "85 · Excelente", TEAL),
    ]),
    ("✍️", "Modo Manual — Transacciones", "Implementado", TEAL, [
        "Añadir gasto/ingreso con tipo, categoría y fecha",
        "Filtros: por tipo (ingreso/gasto), categoría y búsqueda por texto",
        "Ordenar por fecha, monto o categoría (asc/desc)",
        "Editar transacciones existentes",
        "Eliminar con confirmación visual",
        "Resumen: total ingresos, gastos y número de transacciones",
    ], [
        ("Total Ingresos", "+$5,200.00", GREEN),
        ("Total Gastos", "-$2,485.50", RED),
        ("Transacciones", "24 registros", None),
        ("Alquiler Mayo", "-$900.00", RED),
        ("Nómina Abril", "+$3,200.00", GREEN),
        ("Mercadona", "-$156.30", RED),
    ]),
    ("🤖", "Modo IA — Optimizador Inteligente", "Implementado", TEAL, [
        "Plan mensual con distribución por categorías según perfil de gasto",
        "Personalidades: Balanceada, Ahorro Agresivo, Flexible",
        "Chat financiero: preguntas en lenguaje natural sobre tus finanzas",
        "Predicciones: \"Si reduzco un 30% en ocio, ¿cuánto ahorro?\"",
        "Comparativa: gasto actual vs recomendación de la IA",
        "Metas inteligentes con plan de ahorro personalizado",
    ], [
        ("Ingreso Mensual", "$3,000.00", None),
        ("Ahorro Recomendado", "$360.00", GREEN),
        ("Tasa de Ahorro", "12.0%", GREEN),
        ("Vivienda (recom.)", "$840 · 28%", PURPLE),
        ("Comida (recom.)", "$420 · 14%", RGBColor(0xf5, 0x9e, 0x0b)),
        ("Ahorro (recom.)", "$360 · 12%", GREEN),
    ]),
    ("📊", "Presupuestos", "Implementado", TEAL, [
        "Presupuesto mensual por categoría (vivienda, comida, ocio...)",
        "Barras de progreso con color según estado (verde/naranja/rojo)",
        "Crear nuevos presupuestos para categorías sin asignar",
        "Editar límites directamente desde la tarjeta",
        "Resetear gastos mensualmente con un clic",
        "Alerta de presupuestos en riesgo (porcentaje cercano al límite)",
    ], [
        ("Presupuesto Total", "$2,700.00", None),
        ("Vivienda", "$900 / $1,000 · 90%", RGBColor(0xf5, 0x9e, 0x0b)),
        ("Alimentación", "$230 / $500 · 46%", GREEN),
        ("Transporte", "$60 / $200 · 30%", GREEN),
        ("Ocio", "$140 / $150 · 93%", RED),
        ("Saludables", "3 de 5", GREEN),
    ]),
    ("📈", "Estadísticas e Informes", "A mejorar", PURPLE, [
        "Gráfico de tendencia semanal (ingresos vs gastos)",
        "Distribución por categorías en gráfico donut",
        "Comparativa presupuesto vs real con barras agrupadas",
        "Evolución mensual de ahorro acumulado",
        "Top gastos del período seleccionado",
        "Exportar informe en PDF con datos del mes",
    ], [
        ("Ingresos del Mes", "$3,200.00", GREEN),
        ("Gastos del Mes", "$1,485.00", RED),
        ("Ahorro del Mes", "$1,715.00", GREEN),
        ("Top Gasto", "Vivienda · $900", RED),
        ("Categorías con gasto", "6 activas", None),
    ]),
    ("📸", "Escanear Ticket", "Por implementar", PURPLE, [
        "OCR inteligente reconoce texto del ticket",
        "Categorización automática de cada producto",
        "Detección de total, fecha, tienda y método de pago",
        "Vista previa antes de guardar",
        "Corrección manual si la IA se equivoca",
        "Ticket de 20 productos categorizado en segundos",
    ], [
        ("Ticket escaneado", "Mercadona · 28/04/26", None),
        ("Base imponible", "$32.47", None),
        ("IVA", "$3.75", None),
        ("Total detectado", "$36.22", RED),
        ("Categoría asignada", "Alimentación", TEAL),
    ]),
    ("🎯", "Metas de Ahorro", "Por implementar", PURPLE, [
        "Crear meta: nombre, monto objetivo, fecha límite",
        "Plan de ahorro automático generado por IA",
        "Seguimiento visual con progreso hacia la meta",
        "Múltiples metas simultáneas (viaje, coche, fondo emergencia)",
        "Notificaciones de hitos alcanzados",
        "Simulador: \"Si ahorro X más, llego Y semanas antes\"",
    ], [
        ("Viaje a Japón", "$2,250 / $5,000 · 45%", TEAL),
        ("Fondo Coche", "$1,100 / $5,000 · 22%", RGBColor(0xf5, 0x9e, 0x0b)),
        ("Fondo Emergencia", "$3,000 / $6,000 · 50%", GREEN),
        ("Próxima meta", "Viaje · Dic 2026", None),
    ]),
    ("👥", "Finanzas Compartidas", "Por implementar", PURPLE, [
        "Grupos: crear grupo con miembros e invitarlos",
        "Gastos compartidos: dividir equitativamente o personalizado",
        "Liquidación de deudas: \"A le debe X a B\"",
        "Presupuesto grupal con límites por categoría",
        "Historial de gastos del grupo filtrable",
        "Notificaciones cuando alguien añade un gasto compartido",
    ], [
        ("Gasto del mes (grupo)", "$1,240.50", None),
        ("Ana pagó", "$480.00", TEAL),
        ("Bob pagó", "$320.00", TEAL),
        ("Carlos pagó", "$440.00", TEAL),
        ("Deuda pendiente", "Ana debe $15 a Bob", RED),
    ]),
    ("📅", "Calendario Financiero", "Por implementar", PURPLE, [
        "Vista mensual con gastos/ingresos por día",
        "Suscripciones recurrentes marcadas automáticamente",
        "Recordatorios de pagos próximos",
        "Saldo proyectado a fin de mes",
        "Días sin gasto destacados positivamente",
        "Click en día para ver detalle de transacciones",
    ], [
        ("Proyección fin de mes", "+$2,847.50", GREEN),
        ("Próximo pago", "Netflix · -$15.99 · 3 jun", RED),
        ("Días sin gasto", "3 this month", GREEN),
        ("Suscripciones activas", "4 · $67.96/mes", None),
    ]),
    ("🏦", "Conexión Bancaria", "Simulado", TEAL, [
        "Selección de banco entre múltiples entidades",
        "Conexión segura con cifrado simulado TLS 1.3",
        "Importación automática de transacciones de últimos 3 meses",
        "Detección de nómina para ajustar ingreso mensual",
        "Sincronización programable (diaria/semanal)",
        "Desconexión segura en un clic",
    ], [
        ("Banco conectado", "ING · ES12 **** 3456", TEAL),
        ("Última sincronización", "Hoy 09:30", GREEN),
        ("Transacciones importadas", "24 en 3 meses", None),
        ("Nómina detectada", "$3,000.00/mes", GREEN),
    ]),
    ("⚙️", "Ajustes y Configuración", "Implementado", TEAL, [
        "Información financiera: ingreso mensual, moneda (EUR/USD/GBP/MXN), período",
        "Personalidad de IA: Balanceada / Ahorro Agresivo / Flexible",
        "Apariencia: tema claro/oscuro con toggle",
        "Datos: exportar/importar JSON, reseteo completo",
        "Perfil: nombre, email, avatar, cambio de contraseña",
        "Notificaciones: alertas de presupuesto, recordatorios, informes semanales",
    ], [
        ("Ingreso mensual", "$3,000.00", None),
        ("Moneda", "USD · Dólar", None),
        ("Personalidad IA", "Balanceada", TEAL),
        ("Modo oscuro", "Activado", None),
        ("Notificaciones", "Alertas ON", GREEN),
    ]),
]

# ════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════

# ─── COVER PAGE ───
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BudgetIQ")
r.font.size = Pt(40)
r.font.bold = True
r.font.color.rgb = TEAL

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Documentación Visual de Features")
r2.font.size = Pt(16)
r2.font.color.rgb = GRAY

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(8)
r3 = p3.add_run("11 funcionalidades documentadas con descripción y datos de ejemplo")
r3.font.size = Pt(10)
r3.font.color.rgb = MUTED

# Divider
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(16)
r4 = p4.add_run("─" * 50)
r4.font.color.rgb = MUTED
r4.font.size = Pt(8)

# Legend
doc.add_paragraph()
p_leg = doc.add_paragraph()
p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_leg1 = p_leg.add_run("✅ Implementado  ")
r_leg1.font.size = Pt(9)
r_leg1.font.bold = True
r_leg1.font.color.rgb = GREEN
r_leg2 = p_leg.add_run("  🚀 Por implementar  ")
r_leg2.font.size = Pt(9)
r_leg2.font.bold = True
r_leg2.font.color.rgb = PURPLE
r_leg3 = p_leg.add_run("  ⚡ A mejorar")
r_leg3.font.size = Pt(9)
r_leg3.font.bold = True
r_leg3.font.color.rgb = RGBColor(0xf5, 0x9e, 0x0b)

doc.add_page_break()

# ─── TABLE OF CONTENTS ───
p_toc = doc.add_paragraph()
r_toc = p_toc.add_run("ÍNDICE DE FUNCIONALIDADES")
r_toc.font.size = Pt(14)
r_toc.font.bold = True
r_toc.font.color.rgb = DARK

toc_table = doc.add_table(rows=1, cols=4)
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
# Header
for j, h_text in enumerate(["#", "Feature", "Estado", ""]):
    cell = toc_table.rows[0].cells[j]
    p = cell.paragraphs[0]
    r = p.add_run(h_text)
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = WHITE
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0D9488"/>')
    p._element.get_or_add_pPr().append(shading)

for i, (icon, title, status, color, _, _) in enumerate(FEATURES):
    row = toc_table.add_row()
    vals = [str(i+1), f"{icon} {title}", status, ""]
    for j, v in enumerate(vals):
        cell = row.cells[j].paragraphs[0]
        r = cell.add_run(v)
        r.font.size = Pt(8)
        if j == 2:
            r.font.bold = True
            r.font.color.rgb = color
        elif j == 1:
            r.font.bold = True
            r.font.color.rgb = DARK
        else:
            r.font.color.rgb = GRAY

doc.add_page_break()

# ─── EACH FEATURE ───
for i, (icon, title, status, status_color, bullets, mockup_data) in enumerate(FEATURES):
    add_section_title(doc, icon, title, f"  {status}  ", status_color)

    # Description bullets
    for b in bullets:
        colon_idx = b.find(": ")
        if colon_idx > 0 and colon_idx < 50:
            bold_part = b[:colon_idx+2]
            rest = b[colon_idx+2:]
            add_bullet(doc, rest, bold_part)
        else:
            add_bullet(doc, b)

    doc.add_paragraph()

    # Phone mockup as table
    mock_title = doc.add_paragraph()
    mt_r = mock_title.add_run(f"📱 Vista en la app — {title}")
    mt_r.font.size = Pt(9)
    mt_r.font.bold = True
    mt_r.font.color.rgb = MUTED

    add_phone_mockup_table(doc, mockup_data)

    if i < len(FEATURES) - 1:
        doc.add_page_break()

# ─── FOOTER ───
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run("─" * 50)
r_f.font.size = Pt(8)
r_f.font.color.rgb = MUTED

p_last = doc.add_paragraph()
p_last.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_last.paragraph_format.space_before = Pt(8)
r_last = p_last.add_run("BudgetIQ · Documentación de Producto · 2026")
r_last.font.size = Pt(9)
r_last.font.color.rgb = GRAY

# ─── SAVE ───
output_path = "C:\\Users\\Asus\\budgetiq\\BudgetIQ_Features.docx"
doc.save(output_path)
print(f"Documento guardado en: {output_path}")
